from __future__ import annotations

import hashlib
import json
import logging
import os
import re
import shutil
import subprocess
import sys
import tempfile
import unicodedata
import uuid
from collections import Counter
from pathlib import Path
from typing import Any, Iterable

import faiss
import fitz
import httpx
import numpy as np
from docx import Document
from openpyxl import load_workbook
from sentence_transformers import SentenceTransformer

from backend.app.rag.models import PageDocument, TextChunk


logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS = {".pdf", ".md", ".txt", ".docx", ".xlsx", ".json"}
INGESTION_MANIFEST = ".ingestion_manifest.json"
EXTRACTION_CACHE_VERSION = "2"
AD_NOISE = (
    "扫码关注",
    "微信公众号",
    "关注公众号",
    "购买正版",
    "资源下载",
    "学兔兔",
    "bzfxw",
    "广告",
)
TAG_KEYWORDS = (
    "电路模型",
    "参考方向",
    "关联参考方向",
    "欧姆定律",
    "基尔霍夫电流定律",
    "基尔霍夫电压定律",
    "KCL",
    "KVL",
    "节点电压法",
    "回路电流法",
    "网孔电流法",
    "叠加定理",
    "戴维南定理",
    "诺顿定理",
    "最大功率传输定理",
    "含受控源电路",
    "一阶电路",
    "RC电路",
    "RL电路",
    "换路定则",
    "三要素法",
    "时间常数",
    "零输入响应",
    "零状态响应",
    "全响应",
    "二阶电路",
    "正弦稳态",
    "相量法",
    "复阻抗",
    "相位差",
    "有功功率",
    "无功功率",
    "视在功率",
    "复功率",
    "功率因数",
    "RLC谐振",
    "串联谐振",
    "并联谐振",
    "频率响应",
    "传递函数",
    "拉普拉斯变换",
    "二端口网络",
    "运算放大器",
    "本征半导体",
    "N型半导体",
    "P型半导体",
    "PN结",
    "二极管",
    "稳压二极管",
    "晶体管",
    "场效应管",
    "静态工作点",
    "伏安特性",
    "单向导电性",
    "反向击穿",
    "基尔霍夫定律",
)
SECTION_CONCEPT_MARKERS = (
    "电路", "电压", "电流", "电阻", "电容", "电感", "功率", "频率", "信号",
    "半导体", "PN结", "二极管", "晶体管", "场效应管", "MOS", "器件", "集成",
    "放大", "反馈", "振荡", "滤波", "比较器", "运算", "整流", "稳压", "偏置",
    "模型", "特性", "参数", "响应", "网络", "相量", "阻抗", "谐振", "工作点",
)

DOCUMENT_TYPE_TOKENS = {
    "exam": ("试卷", "考试", "期末", "期中", "月考", "真题", "模拟卷", "测试卷", "练习卷"),
    "question_bank": ("题库", "习题集", "练习题", "习题答案"),
    "textbook": ("教材", "教程", "讲义", "教科书", "课程", "基础", "原理"),
}


def _normalize_line(line: str) -> str:
    line = unicodedata.normalize("NFKC", line).replace("\u200b", "")
    line = re.sub(r"https?\s*:\s*[/\\]+\s*\S+", "", line, flags=re.I)
    line = re.sub(r"www\s*\.\s*\S+", "", line, flags=re.I)
    line = re.sub(r"(?<=[A-Za-z])\s+(?=[A-Za-z0-9])", "", line)
    line = re.sub(r"(?<=[0-9])\s+(?=[A-Za-z])", "", line)
    line = re.sub(r"(?<=[A-Za-z])\s+(?=[0-9])", "", line)
    line = re.sub(r"P\s*N\s*结", "PN结", line, flags=re.I)
    line = re.sub(r"([NP])\s*型", r"\1型", line, flags=re.I)
    line = re.sub(r"M\s*O\s*S", "MOS", line, flags=re.I)
    line = re.sub(r"[ \t]+", " ", line).strip()
    return line


def _edge_noise(raw_pages: list[str]) -> set[str]:
    candidates: Counter[str] = Counter()
    for text in raw_pages:
        lines = [_normalize_line(line) for line in text.splitlines() if _normalize_line(line)]
        for line in lines[:2] + lines[-2:]:
            if len(line) <= 80:
                candidates[line] += 1
    threshold = max(3, int(len(raw_pages) * 0.08))
    return {line for line, count in candidates.items() if count >= threshold}


def clean_page_text(text: str, repeated_noise: set[str] | None = None) -> str:
    repeated_noise = repeated_noise or set()
    cleaned_lines: list[str] = []
    for raw_line in text.splitlines():
        line = _normalize_line(raw_line)
        if not line or line in repeated_noise:
            continue
        if re.fullmatch(r"[-—·•\s]*\d{1,4}[-—·•\s]*", line):
            continue
        if any(noise in line for noise in AD_NOISE):
            continue
        if sum(char == "�" for char in line) > 1:
            continue
        cleaned_lines.append(line)

    paragraphs: list[str] = []
    buffer = ""
    heading_pattern = re.compile(
        r"^(?:第[一二三四五六七八九十百0-9]+章|\d+(?:\.\d+){0,3}\s+|本章小结|自测题|习题)"
    )
    for line in cleaned_lines:
        is_heading = bool(heading_pattern.match(line)) and len(line) < 70
        if is_heading:
            if buffer:
                paragraphs.append(buffer)
                buffer = ""
            paragraphs.append(line)
            continue
        buffer += line
        if line.endswith(("。", "！", "？", ":", "；")) or len(buffer) >= 260:
            paragraphs.append(buffer)
            buffer = ""
    if buffer:
        paragraphs.append(buffer)
    return "\n\n".join(paragraphs)


def _is_chapter_title(title: str) -> bool:
    return bool(re.match(r"^第[一二三四五六七八九十百0-9]+章", title.replace(" ", "")))


def _pdf_page_range(
    document: fitz.Document, chapter_limit: int | None
) -> tuple[int, int, list[list[Any]]]:
    toc = [item for item in document.get_toc(simple=True) if len(item) >= 3]
    chapters = [
        (str(title).strip(), int(page))
        for level, title, page in toc
        if level == 1 and _is_chapter_title(str(title).strip())
    ]
    if not chapters:
        return 1, document.page_count, toc
    start_page = chapters[0][1]
    end_page = (
        chapters[chapter_limit][1] - 1
        if chapter_limit and len(chapters) > chapter_limit
        else document.page_count
    )
    return start_page, end_page, toc


def _pdf_hierarchy(
    toc: list[list[Any]], page_number: int, fallback: str
) -> tuple[str, str]:
    chapter = ""
    section = ""
    for level, title, toc_page, *_ in toc:
        if int(toc_page) > page_number:
            break
        normalized = _normalize_line(re.sub(r"\s+", " ", str(title)).strip())
        if level == 1 and _is_chapter_title(normalized):
            chapter = normalized
            section = ""
        elif level >= 2:
            section = normalized
    return chapter or fallback, section or chapter or fallback


def _infer_ocr_hierarchy(
    text: str, fallback: str, chapter: str, section: str
) -> tuple[str, str]:
    """Recover printed chapter/section headings when a scan has no PDF outline."""
    lines = [_normalize_line(line) for line in text.splitlines()[:16]]
    lines = [line for line in lines if 3 <= len(line) <= 90]
    current_chapter = chapter or fallback
    current_section = section or current_chapter
    for line in lines:
        compact = line.replace(" ", "")
        looks_like_sentence = bool(
            re.match(
                r"^第[一二三四五六七八九十百0-9]+章(?:中|已|我们|介绍|将|讨论)",
                compact,
            )
        )
        if _is_chapter_title(compact) and not looks_like_sentence and len(compact) <= 36:
            current_chapter = line
            current_section = line
            continue
        section_match = re.match(r"^(\d+)(?:\.\d+){1,3}\s*[^\d\W].+", line)
        if section_match or re.match(
            r"^[一二三四五六七八九十]+[、.]\s*.+", line
        ):
            current_section = line
            if section_match:
                major = int(section_match.group(1))
                numerals = "零一二三四五六七八九十"
                inferred_prefix = f"第{numerals[major]}章" if 0 < major < len(numerals) else f"第{major}章"
                if not current_chapter.startswith(inferred_prefix):
                    current_chapter = inferred_prefix
    return current_chapter, current_section


def extract_pdf(
    path: Path,
    chapter_limit: int | None = None,
    *,
    doc_type: str = "textbook",
) -> list[PageDocument]:
    document = fitz.open(path)
    start_page, end_page, toc = _pdf_page_range(document, chapter_limit)

    raw_pages = [document[index - 1].get_text("text") for index in range(start_page, end_page + 1)]
    repeated_noise = _edge_noise(raw_pages)
    page_docs: list[PageDocument] = []
    for page_number, raw_text in zip(range(start_page, end_page + 1), raw_pages):
        current_chapter, current_section = _pdf_hierarchy(toc, page_number, path.stem)
        text = clean_page_text(raw_text, repeated_noise)
        if len(text) < 30:
            continue
        page_docs.append(
            PageDocument(
                text=text,
                source=path.name,
                page=page_number,
                chapter=current_chapter,
                section=current_section,
                doc_type=doc_type,
            )
        )
    document.close()
    return page_docs


def _extract_pdf_macos_vision(
    path: Path, chapter_limit: int | None = None
) -> list[PageDocument]:
    """Optional zero-install OCR fallback for macOS.

    The main parser remains portable: production deployments should configure
    PDF_EXTRACTOR_URL.  This adapter only activates automatically on macOS when
    a scan has no usable text layer.
    """
    if sys.platform != "darwin" or os.getenv("PDF_LOCAL_OCR", "auto").lower() == "off":
        return []
    xcrun = shutil.which("xcrun")
    source = Path(__file__).resolve().parents[3] / "scripts" / "macos_vision_ocr.m"
    if not xcrun or not source.exists():
        return []

    raw_pages: list[tuple[int, str]] = []
    with tempfile.TemporaryDirectory(prefix="circuitmind-ocr-") as temporary:
        workspace = Path(temporary)
        executable = workspace / "macos_vision_ocr"
        module_cache = workspace / "clang-module-cache"
        module_cache.mkdir()
        environment = {**os.environ, "CLANG_MODULE_CACHE_PATH": str(module_cache)}
        subprocess.run(
            [
                xcrun,
                "clang",
                "-fobjc-arc",
                "-fblocks",
                str(source),
                "-framework",
                "Foundation",
                "-framework",
                "AppKit",
                "-framework",
                "Vision",
                "-o",
                str(executable),
            ],
            check=True,
            capture_output=True,
            text=True,
            timeout=120,
            env=environment,
        )
        with fitz.open(path) as document:
            start_page, end_page, toc = _pdf_page_range(document, chapter_limit)
            page_numbers = list(range(start_page, end_page + 1))
            for offset in range(0, len(page_numbers), 8):
                batch = page_numbers[offset : offset + 8]
                image_paths: list[Path] = []
                for page_number in batch:
                    image_path = workspace / f"page-{page_number}.png"
                    pixmap = document[page_number - 1].get_pixmap(
                        matrix=fitz.Matrix(2, 2), colorspace=fitz.csRGB, alpha=False
                    )
                    pixmap.save(image_path)
                    image_paths.append(image_path)
                completed = subprocess.run(
                    [str(executable), *(str(item) for item in image_paths)],
                    check=True,
                    capture_output=True,
                    text=True,
                    timeout=300,
                    env=environment,
                )
                batch_results = json.loads(completed.stdout)
                for index, item in enumerate(batch_results):
                    raw_pages.append((batch[index], str(item.get("text", ""))))
                for image_path in image_paths:
                    image_path.unlink(missing_ok=True)
                logger.info(
                    "OCR %s: %s/%s pages",
                    path.name,
                    min(offset + len(batch), len(page_numbers)),
                    len(page_numbers),
                )

            repeated_noise = _edge_noise([text for _, text in raw_pages])
            documents: list[PageDocument] = []
            inferred_chapter = path.stem
            inferred_section = path.stem
            for page_number, raw_text in raw_pages:
                text = clean_page_text(raw_text, repeated_noise)
                if len(text) < 30:
                    continue
                chapter, section = _pdf_hierarchy(toc, page_number, path.stem)
                if chapter == path.stem and section == path.stem:
                    inferred_chapter, inferred_section = _infer_ocr_hierarchy(
                        text, path.stem, inferred_chapter, inferred_section
                    )
                    chapter, section = inferred_chapter, inferred_section
                documents.append(
                    PageDocument(
                        text=text,
                        source=path.name,
                        page=page_number,
                        chapter=chapter,
                        section=section,
                    )
                )
            return documents


def extract_pdf_adaptive(
    path: Path,
    chapter_limit: int | None = None,
    *,
    extractor_url: str = "",
) -> tuple[list[PageDocument], str, list[str]]:
    """Use an optional layout/OCR service, then fall back to local PyMuPDF.

    The service boundary is deliberately generic so PDF-Extract-Kit, Docling or
    another parser can sit behind the same adapter. It accepts multipart `file`
    and returns `{pages: [{page, text|markdown, chapter?, section?}]}`.
    """
    warnings: list[str] = []
    if extractor_url:
        try:
            with path.open("rb") as handle, httpx.Client(timeout=180.0, trust_env=False) as client:
                response = client.post(
                    extractor_url,
                    files={"file": (path.name, handle, "application/pdf")},
                    data={"output_format": "pages", "preserve_formulas": "true"},
                )
                response.raise_for_status()
            payload = response.json()
            pages = payload.get("pages", payload.get("data", {}).get("pages", []))
            documents: list[PageDocument] = []
            for index, item in enumerate(pages, 1):
                if not isinstance(item, dict):
                    continue
                # Layout-aware adapters often name their output `markdown` so
                # formulas and tables survive as LaTeX/Markdown instead of OCR
                # character soup. Plain `text` remains fully compatible.
                content = str(item.get("markdown") or item.get("text") or "")
                if len(content.strip()) < 30:
                    continue
                documents.append(
                    PageDocument(
                        text=clean_page_text(content),
                        source=path.name,
                        page=int(item.get("page") or item.get("page_number") or index),
                        chapter=str(item.get("chapter") or path.stem),
                        section=str(item.get("section") or item.get("chapter") or path.stem),
                    )
                )
            if documents:
                parser = str(payload.get("parser") or "external-layout-ocr")
                return documents, parser, warnings
            warnings.append("外部布局/OCR解析器未返回有效页面，已回退 PyMuPDF")
        except Exception as exc:
            warnings.append(f"外部布局/OCR解析失败，已回退 PyMuPDF：{exc}")

    documents = extract_pdf(path, chapter_limit)
    with fitz.open(path) as pdf:
        start_page, end_page, _ = _pdf_page_range(pdf, chapter_limit)
        expected_pages = end_page - start_page + 1
    if len(documents) < max(1, int(expected_pages * 0.2)):
        warnings.append("原生文本覆盖率较低，检测为扫描版 PDF")
        try:
            ocr_documents = _extract_pdf_macos_vision(path, chapter_limit)
        except Exception as exc:
            warnings.append(f"本地 OCR 解析失败：{exc}")
            ocr_documents = []
        if ocr_documents:
            warnings.append("已自动使用 macOS Vision OCR；跨平台部署建议配置 PDF_EXTRACTOR_URL")
            return ocr_documents, "macos-vision-ocr", warnings
        warnings.append("请配置 PDF_EXTRACTOR_URL 或安装可用的本地 OCR 解析器")
    return documents, "pymupdf", warnings


def extract_markdown_or_text(path: Path) -> list[PageDocument]:
    text = path.read_text(encoding="utf-8", errors="ignore")
    chapter = path.stem
    section = path.stem
    documents: list[PageDocument] = []
    page = 1
    buffer: list[str] = []

    def flush() -> None:
        nonlocal buffer
        cleaned = clean_page_text("\n".join(buffer))
        if cleaned:
            documents.append(PageDocument(cleaned, path.name, page, chapter, section))
        buffer = []

    for line in text.splitlines():
        heading = re.match(r"^(#{1,3})\s+(.+)$", line.strip())
        if heading:
            flush()
            title = heading.group(2).strip()
            if len(heading.group(1)) == 1:
                chapter = title
            section = title
        elif line.strip() == "\f":
            flush()
            page += 1
        else:
            buffer.append(line)
    flush()
    return documents


def extract_docx(path: Path) -> list[PageDocument]:
    document = Document(path)
    chapter = path.stem
    section = path.stem
    blocks: list[PageDocument] = []
    buffer: list[str] = []

    def flush() -> None:
        nonlocal buffer
        text = clean_page_text("\n".join(buffer))
        if text:
            blocks.append(PageDocument(text, path.name, 1, chapter, section))
        buffer = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style = (paragraph.style.name or "").lower()
        if "heading" in style or "标题" in style:
            flush()
            section = text
            if style.endswith("1"):
                chapter = text
        else:
            buffer.append(text)
    flush()
    return blocks


def _split_tags(value: Any) -> list[str]:
    if isinstance(value, list):
        return [str(item).strip() for item in value if str(item).strip()]
    return [item.strip() for item in re.split(r"[,，、;；|]", str(value or "")) if item.strip()]


QUESTION_HEADER_ALIASES = {
    "question_id": ("题号", "question_id", "id"),
    "question_text": ("题目文本", "题目", "question_text"),
    "knowledge_tags": ("知识点标签", "知识点", "knowledge_tags"),
    "standard_answer": ("标准答案", "答案", "standard_answer"),
    "common_mistakes": ("易错点", "common_mistakes"),
    "difficulty": ("难度", "difficulty"),
    "question_type": ("题型", "question_type"),
    "solution_steps": ("解题步骤", "解析", "solution_steps"),
}


def extract_question_xlsx(path: Path) -> list[dict[str, Any]]:
    workbook = load_workbook(path, read_only=True, data_only=True)
    sheet = workbook.active
    rows = list(sheet.iter_rows(values_only=True))
    if not rows:
        return []
    header_index = next(
        (
            index
            for index, row in enumerate(rows[:12])
            if "题号" in [str(value or "").strip() for value in row]
            and "题目文本" in [str(value or "").strip() for value in row]
        ),
        0,
    )
    headers = [str(value or "").strip() for value in rows[header_index]]
    mapping: dict[str, int] = {}
    for field, aliases in QUESTION_HEADER_ALIASES.items():
        for alias in aliases:
            if alias in headers:
                mapping[field] = headers.index(alias)
                break
    required = {"question_id", "question_text", "knowledge_tags", "standard_answer", "common_mistakes"}
    if not required.issubset(mapping):
        workbook.close()
        return []
    questions: list[dict[str, Any]] = []
    for row in rows[header_index + 1 :]:
        if not row or not row[mapping["question_text"]]:
            continue
        item = {
            field: (row[index] if index < len(row) else "")
            for field, index in mapping.items()
        }
        item["question_id"] = str(item["question_id"])
        item["question_text"] = str(item["question_text"]).strip()
        item["knowledge_tags"] = _split_tags(item["knowledge_tags"])
        item["standard_answer"] = str(item["standard_answer"] or "").strip()
        item["common_mistakes"] = str(item["common_mistakes"] or "").strip()
        item["difficulty"] = str(item.get("difficulty") or "基础").strip()
        item["question_type"] = str(item.get("question_type") or "综合题").strip()
        item["solution_steps"] = str(item.get("solution_steps") or "").strip()
        item["source"] = path.name
        questions.append(item)
    workbook.close()
    return questions


def extract_question_json(path: Path) -> list[dict[str, Any]]:
    try:
        data = json.loads(path.read_text(encoding="utf-8"))
    except (json.JSONDecodeError, UnicodeError):
        return []
    if isinstance(data, dict):
        data = data.get("questions", [])
    if not isinstance(data, list):
        return []
    questions = []
    for index, item in enumerate(data, 1):
        if not isinstance(item, dict) or not item.get("question_text"):
            continue
        normalized = dict(item)
        normalized.setdefault("question_id", f"JSON-{index:03d}")
        normalized["knowledge_tags"] = _split_tags(normalized.get("knowledge_tags"))
        normalized.setdefault("standard_answer", "")
        normalized.setdefault("common_mistakes", "")
        normalized.setdefault("difficulty", "基础")
        normalized.setdefault("question_type", "综合题")
        normalized.setdefault("solution_steps", "")
        normalized["source"] = path.name
        questions.append(normalized)
    return questions


def _load_ingestion_manifest(resources_dir: Path) -> dict[str, str]:
    path = resources_dir / INGESTION_MANIFEST
    if not path.exists():
        return {}
    try:
        value = json.loads(path.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError):
        return {}
    files = value.get("files", value) if isinstance(value, dict) else {}
    return {
        str(name): str(metadata.get("document_type", metadata) if isinstance(metadata, dict) else metadata)
        for name, metadata in files.items()
        if str(metadata.get("document_type", metadata) if isinstance(metadata, dict) else metadata)
        in {"auto", "textbook", "exam", "question_bank", "notes"}
    }


def list_source_files(resources_dir: Path) -> list[Path]:
    """Return user content only; internal dotfiles never become course sources."""
    return [
        path
        for path in sorted(resources_dir.iterdir())
        if path.is_file()
        and not path.name.startswith(".")
        and path.name != INGESTION_MANIFEST
        and path.suffix.lower() in SUPPORTED_EXTENSIONS
    ]


def _pdf_cache_path(
    cache_dir: Path,
    path: Path,
    chapter_limit: int | None,
    extractor_url: str,
) -> Path:
    stat = path.stat()
    signature = "|".join(
        (
            EXTRACTION_CACHE_VERSION,
            str(path.resolve()),
            str(stat.st_size),
            str(stat.st_mtime_ns),
            str(chapter_limit),
            extractor_url,
            os.getenv("PDF_LOCAL_OCR", "auto"),
        )
    )
    return cache_dir / f"{hashlib.sha1(signature.encode('utf-8')).hexdigest()}.json"


def _load_cached_pdf(cache_path: Path) -> tuple[list[PageDocument], str, list[str]] | None:
    try:
        payload = json.loads(cache_path.read_text(encoding="utf-8"))
        if payload.get("schema_version") != EXTRACTION_CACHE_VERSION:
            return None
        documents = [PageDocument(**item) for item in payload.get("documents", [])]
        if not documents:
            return None
        return documents, str(payload.get("parser", "cache")), list(payload.get("warnings", []))
    except (OSError, TypeError, ValueError, json.JSONDecodeError):
        return None


def _write_cached_pdf(
    cache_path: Path,
    documents: list[PageDocument],
    parser: str,
    warnings: list[str],
) -> None:
    if not documents:
        return
    cache_path.parent.mkdir(parents=True, exist_ok=True)
    temporary = cache_path.with_name(f".{cache_path.name}.{uuid.uuid4().hex[:8]}.tmp")
    temporary.write_text(
        json.dumps(
            {
                "schema_version": EXTRACTION_CACHE_VERSION,
                "parser": parser,
                "warnings": warnings,
                "documents": [document.__dict__ for document in documents],
            },
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )
    temporary.replace(cache_path)


def classify_document(path: Path, documents: list[PageDocument], declared: str = "auto") -> str:
    """Classify a source without coupling ingestion to a particular PDF parser."""
    if declared != "auto":
        return declared
    if path.suffix.lower() in {".xlsx", ".json"}:
        return "question_bank"
    filename = path.stem.lower().replace(" ", "")
    for document_type in ("exam", "question_bank", "textbook"):
        if any(token in filename for token in DOCUMENT_TYPE_TOKENS[document_type]):
            return document_type
    sample = "\n".join(document.text for document in documents[:12])
    numbered = len(re.findall(r"(?:^|\n)\s*(?:\d{1,3}[.、．)]|[一二三四五六七八九十]+[、.])\s*", sample))
    question_words = len(re.findall(r"(?:选择题|填空题|计算题|简答题|证明题|本题\s*\d+\s*分)", sample))
    return "exam" if numbered >= 8 and question_words >= 1 else "textbook"


def extract_questions_from_documents(
    documents: list[PageDocument], source: str
) -> list[dict[str, Any]]:
    """Extract question stems from exam-like documents with stable source anchors.

    This lightweight parser intentionally produces reviewable candidates. Layout/OCR
    engines can be added before this stage without changing the downstream schema.
    """
    questions: list[dict[str, Any]] = []
    marker = re.compile(
        r"(?m)^(?:\s*)(?P<label>(?:第\s*)?\d{1,3}\s*[.、．）)]|[一二三四五六七八九十]+[、.])\s*"
    )
    for document in documents:
        matches = list(marker.finditer(document.text))
        for index, match in enumerate(matches):
            end = matches[index + 1].start() if index + 1 < len(matches) else len(document.text)
            stem = document.text[match.start() : end].strip()
            stem = re.split(r"\n\n(?:参考答案|答案|解析)[:：]?", stem, maxsplit=1)[0].strip()
            if not 12 <= len(stem) <= 5000:
                continue
            digest = hashlib.sha1(f"{source}|{document.page}|{stem[:240]}".encode("utf-8")).hexdigest()[:12]
            tags = _knowledge_tags(stem, document.section)
            questions.append(
                {
                    "question_id": f"AUTO-{digest}",
                    "question_text": stem,
                    "knowledge_tags": tags,
                    "standard_answer": "",
                    "common_mistakes": "",
                    "difficulty": "待标注",
                    "question_type": "自动抽取题",
                    "solution_steps": "",
                    "source": source,
                    "source_page": document.page,
                    "extraction": "rule_candidate",
                }
            )
    return questions


def _knowledge_tags(text: str, section: str) -> list[str]:
    tags = [keyword for keyword in TAG_KEYWORDS if keyword.lower() in text.lower()]
    normalized_section = re.sub(r"^\d+(?:\.\d+)*\s*", "", section).strip()
    normalized_section = re.sub(
        r"^[一二三四五六七八九十]+[、.．]\s*", "", normalized_section
    ).strip()
    section_is_concept = bool(
        2 <= len(normalized_section) <= 22
        and not re.search(r"[,，。！？?:：=;；]", normalized_section)
        and not re.match(r"^(?:第?[一二三四五六七八九十百0-9]+章|本章|习题|思考题)", normalized_section)
        and not any(
            marker in normalized_section
            for marker in ("如图", "见图", "题解", "例题", "已知", "试求", "请判断", "设在", "通过调节", "改变")
        )
        and any(marker.casefold() in normalized_section.casefold() for marker in SECTION_CONCEPT_MARKERS)
    )
    if section_is_concept and normalized_section not in tags:
        tags.insert(0, normalized_section)
    return tags[:8]


def _sentence_pieces(text: str, max_chars: int = 900) -> list[str]:
    paragraphs = [paragraph.strip() for paragraph in text.split("\n\n") if paragraph.strip()]
    pieces: list[str] = []
    for paragraph in paragraphs:
        if len(paragraph) <= max_chars:
            pieces.append(paragraph)
            continue
        sentences = [piece for piece in re.split(r"(?<=[。！？；])", paragraph) if piece]
        current = ""
        for sentence in sentences:
            if current and len(current) + len(sentence) > max_chars:
                pieces.append(current)
                current = ""
            if len(sentence) > max_chars:
                pieces.extend(
                    sentence[index : index + max_chars]
                    for index in range(0, len(sentence), max_chars)
                )
                continue
            current += sentence
        if current:
            pieces.append(current)
    return pieces


def chunk_documents(documents: Iterable[PageDocument], max_chars: int = 900) -> list[TextChunk]:
    chunks: list[TextChunk] = []
    for document in documents:
        pieces = _sentence_pieces(document.text, max_chars=max_chars)
        current: list[str] = []
        current_length = 0

        def flush() -> None:
            nonlocal current, current_length
            if not current:
                return
            text = "\n\n".join(current).strip()
            raw_id = f"{document.source}|{document.page}|{document.section}|{text[:120]}"
            chunk_id = hashlib.sha1(raw_id.encode("utf-8")).hexdigest()[:16]
            chunks.append(
                TextChunk(
                    id=chunk_id,
                    text=text,
                    source=document.source,
                    chapter=document.chapter,
                    section=document.section,
                    page_start=document.page,
                    page_end=document.page,
                    doc_type=document.doc_type,
                    knowledge_tags=_knowledge_tags(text, document.section),
                )
            )
            overlap = text[-140:] if len(text) > 140 else ""
            current = [overlap] if overlap else []
            current_length = len(overlap)

        for piece in pieces:
            if current and current_length + len(piece) > max_chars:
                flush()
            if current and current_length + len(piece) > max_chars:
                current = []
                current_length = 0
            current.append(piece)
            current_length += len(piece)
        flush()
    return chunks


def question_chunks(questions: Iterable[dict[str, Any]]) -> list[TextChunk]:
    chunks: list[TextChunk] = []
    for item in questions:
        tags = _split_tags(item.get("knowledge_tags"))
        text = (
            f"题目：{item.get('question_text', '')}\n"
            f"标准答案：{item.get('standard_answer', '')}\n"
            f"解题步骤：{item.get('solution_steps', '')}\n"
            f"易错点：{item.get('common_mistakes', '')}"
        ).strip()
        question_id = str(item.get("question_id", "Q"))
        chunks.append(
            TextChunk(
                id=f"question-{question_id}",
                text=text,
                source=str(item.get("source", "question_bank.json")),
                chapter="示例题库",
                section="、".join(tags) or "综合",
                page_start=None,
                page_end=None,
                doc_type="question",
                knowledge_tags=tags,
            )
        )
    return chunks


def _write_clean_markdown(path: Path, documents: list[PageDocument], output_dir: Path) -> None:
    output_dir.mkdir(parents=True, exist_ok=True)
    lines = [f"# {path.stem}", "", f"> 清洗来源：{path.name}", ""]
    last_chapter = last_section = ""
    for document in documents:
        if document.chapter and document.chapter != last_chapter:
            lines.extend([f"## {document.chapter}", ""])
            last_chapter = document.chapter
        if document.section and document.section != last_section and document.section != document.chapter:
            lines.extend([f"### {document.section}", ""])
            last_section = document.section
        lines.extend([f"<!-- source={document.source}; page={document.page} -->", document.text, ""])
    (output_dir / f"{path.stem}.clean.md").write_text("\n".join(lines), encoding="utf-8")


def build_knowledge_base(
    resources_dir: Path,
    output_dir: Path,
    embedding_model_path: Path,
    *,
    chapter_limit: int | None = None,
    pdf_extractor_url: str | None = None,
) -> dict[str, Any]:
    resources_dir = resources_dir.resolve()
    output_dir = output_dir.resolve()
    output_dir.mkdir(parents=True, exist_ok=True)
    cleaned_dir = output_dir / "cleaned_documents"
    cleaned_dir.mkdir(parents=True, exist_ok=True)

    documents: list[PageDocument] = []
    questions: list[dict[str, Any]] = []
    source_files = list_source_files(resources_dir)
    declared_types = _load_ingestion_manifest(resources_dir)
    pdf_extractor_url = (pdf_extractor_url if pdf_extractor_url is not None else os.getenv("PDF_EXTRACTOR_URL", "")).strip()
    extraction_cache = output_dir.parent / ".extraction_cache"
    source_manifest: list[dict[str, Any]] = []
    for path in source_files:
        suffix = path.suffix.lower()
        declared_type = declared_types.get(path.name, "auto")
        extracted: list[PageDocument] = []
        extracted_questions: list[dict[str, Any]] = []
        parser_name = suffix.lstrip(".")
        warnings: list[str] = []
        cache_hit = False
        if suffix == ".pdf":
            cache_path = _pdf_cache_path(
                extraction_cache, path, chapter_limit, pdf_extractor_url
            )
            cached = _load_cached_pdf(cache_path)
            if cached:
                extracted, parser_name, warnings = cached
                cache_hit = True
            else:
                extracted, parser_name, warnings = extract_pdf_adaptive(
                    path, chapter_limit, extractor_url=pdf_extractor_url
                )
                _write_cached_pdf(cache_path, extracted, parser_name, warnings)
        elif suffix in {".md", ".txt"}:
            extracted = extract_markdown_or_text(path)
        elif suffix == ".docx":
            extracted = extract_docx(path)
        elif suffix == ".xlsx":
            extracted_questions = extract_question_xlsx(path)
        elif suffix == ".json":
            extracted_questions = extract_question_json(path)

        document_type = classify_document(path, extracted, declared_type)
        for document in extracted:
            document.doc_type = document_type
        if extracted:
            documents.extend(extracted)
            _write_clean_markdown(path, extracted, cleaned_dir)
        if document_type in {"exam", "question_bank"} and extracted:
            extracted_questions.extend(extract_questions_from_documents(extracted, path.name))
        questions.extend(extracted_questions)
        if suffix == ".pdf" and not extracted:
            warnings.append("未提取到文本，可能是扫描版 PDF；请配置布局/OCR解析器后重建")
        source_manifest.append(
            {
                "source": path.name,
                "document_type": document_type,
                "declared_type": declared_type,
                "pages_or_sections": len(extracted),
                "questions": len(extracted_questions),
                "parser": parser_name,
                "cache_hit": cache_hit,
                "warnings": warnings,
            }
        )

    # Course materials and the structured question bank have different trust
    # and retrieval semantics.  Only course chunks enter the RAG vector index;
    # questions stay in question_bank.json and are queried by the quiz workflow.
    chunks = chunk_documents(documents)
    question_items = question_chunks(questions)
    if not chunks:
        raise RuntimeError(f"在 {resources_dir} 中没有提取到可索引内容")

    chunk_path = output_dir / "chunks.jsonl"
    chunk_path.write_text(
        "\n".join(json.dumps(chunk.to_dict(), ensure_ascii=False) for chunk in chunks),
        encoding="utf-8",
    )

    model = SentenceTransformer(str(embedding_model_path), device="cpu")
    embedding_texts = [
        "\n".join(
            filter(
                None,
                [chunk.doc_type, chunk.chapter, chunk.section, " ".join(chunk.knowledge_tags), chunk.text],
            )
        )
        for chunk in chunks
    ]
    embeddings = model.encode(
        embedding_texts,
        batch_size=32,
        show_progress_bar=True,
        normalize_embeddings=True,
        convert_to_numpy=True,
    ).astype(np.float32)
    index = faiss.IndexFlatIP(embeddings.shape[1])
    index.add(embeddings)
    # FAISS' Windows file writer cannot open paths containing Chinese characters.
    # Serialize in memory and let Python handle the Unicode path instead.
    serialized_index = faiss.serialize_index(index)
    (output_dir / "vectors.faiss").write_bytes(serialized_index.tobytes())

    relations: list[dict[str, Any]] = []
    related_by_question: dict[str, list[dict[str, Any]]] = {}
    if question_items:
        question_embeddings = model.encode(
            [
                "\n".join(
                    filter(
                        None,
                        [item.section, " ".join(item.knowledge_tags), item.text],
                    )
                )
                for item in question_items
            ],
            batch_size=32,
            show_progress_bar=False,
            normalize_embeddings=True,
            convert_to_numpy=True,
        ).astype(np.float32)
    else:
        question_embeddings = np.empty((0, embeddings.shape[1]), dtype=np.float32)
    textbook_indices = [
        index for index, chunk in enumerate(chunks) if chunk.doc_type in {"textbook", "notes"}
    ]
    for question_index, chunk in enumerate(question_items):
        if not textbook_indices:
            break
        similarities = embeddings[textbook_indices] @ question_embeddings[question_index]
        shared_tag_sets = [
            set(chunk.knowledge_tags) & set(chunks[target_index].knowledge_tags)
            for target_index in textbook_indices
        ]
        combined_scores = np.asarray(
            [
                float(similarity) + min(0.18, 0.08 * len(shared_tags))
                for similarity, shared_tags in zip(similarities, shared_tag_sets)
            ],
            dtype=np.float32,
        )
        ranked = np.argsort(combined_scores)[::-1][:5]
        question_id = chunk.id.removeprefix("question-")
        for rank in ranked:
            target_index = textbook_indices[int(rank)]
            target = chunks[target_index]
            shared_tags = sorted(shared_tag_sets[int(rank)])
            semantic_score = float(similarities[int(rank)])
            score = float(combined_scores[int(rank)])
            if semantic_score < 0.22 and not shared_tags:
                continue
            relation = {
                "question_id": question_id,
                "question_chunk_id": chunk.id,
                "knowledge_chunk_id": target.id,
                "knowledge_source": target.source,
                "chapter": target.chapter,
                "section": target.section,
                "page": target.page_start,
                "shared_knowledge_tags": shared_tags,
                "score": round(score, 4),
                "semantic_score": round(semantic_score, 4),
                "relation_type": "supported_by",
            }
            relations.append(relation)
            related_by_question.setdefault(question_id, []).append(relation)

    for question in questions:
        question["related_knowledge"] = related_by_question.get(str(question.get("question_id")), [])
    structured_questions = {"schema_version": "1.1", "questions": questions}
    (output_dir / "question_bank.json").write_text(
        json.dumps(structured_questions, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    (output_dir / "knowledge_relations.jsonl").write_text(
        "\n".join(json.dumps(item, ensure_ascii=False) for item in relations), encoding="utf-8"
    )
    (output_dir / "source_manifest.json").write_text(
        json.dumps({"schema_version": "1.0", "sources": source_manifest}, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )

    metadata = {
        "state": "populated",
        "resource_dir": str(resources_dir),
        "embedding_model": str(embedding_model_path),
        "dimension": int(embeddings.shape[1]),
        "documents": len(source_files),
        "indexed_documents": sum(
            1
            for item in source_manifest
            if item["pages_or_sections"] > 0 or item["questions"] > 0
        ),
        "failed_documents": sum(
            1
            for item in source_manifest
            if item["pages_or_sections"] == 0 and item["questions"] == 0
        ),
        "text_pages": len(documents),
        "questions": len(questions),
        "relations": len(relations),
        "chunks": len(chunks),
        "chapter_limit": chapter_limit,
        "sources": [path.name for path in source_files],
        "source_manifest": source_manifest,
    }
    (output_dir / "index_meta.json").write_text(
        json.dumps(metadata, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    logger.info("Knowledge base populated: %s", metadata)
    return metadata
